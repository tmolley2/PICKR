#!/usr/bin/env python3
# -----------------------------------------------------------------------------------
# 0.  Imports ────────────────────────────────────────────────────────
# -----------------------------------------------------------------------------------
import json, re, time, requests, os, gzip, calendar, datetime # Added datetime
import camelot
import pandas as pd
# import glob # Not explicitly used in the final version, can be removed if truly unused
# import textwrap, sys, pprint, itertools # Not explicitly used, can be removed
from pathlib import Path
# from datetime import datetime # datetime is imported from the main datetime module now
from typing import List, Dict, Tuple, Optional
from lxml import etree      # robust XML parser
from tqdm import tqdm # progress-bar
from pebble import ProcessPool
from concurrent.futures import TimeoutError as PebbleTimeoutError
from concurrent.futures import ProcessPoolExecutor, as_completed
import docx
import pdfplumber # <<< Ensure pdfplumber is imported
from multiprocessing import get_context
import signal
from functools import wraps
import shutil # <<< ADD THIS IMPORT FOR DIRECTORY DELETION

# --- Define SCRIPT_ROOT early if needed for TMP_DIR_FOR_SCRIPT ---
# This helps in creating paths relative to the script's location.
SCRIPT_ROOT = Path(__file__).resolve().parent

# --- Configuration: Keywords and Regular Expressions ---
FORWARD_KEYWORDS = ['forward', 'fw', 'fp', 'fwd', 'sense', 'forw']
REVERSE_KEYWORDS = ['reverse', 'rev', 'rp', 'rv', 'antisense', 'revs']
GENE_HEADER_KEYWORDS = ['gene', 'target', 'name', 'symbol', 'oligonucleotide', 'oligos', 'locus']
ALLOWED_DNA_CHARS_FOR_FILTER = "ACTG"


# -----------------------------------------------------------------------------------
# 0.5  Directory and API Configuration ───────────────────────────────────────────
# -----------------------------------------------------------------------------------
DATA_DIR_BASE = Path("data/") # Base data directory
DATA_DIR_BASE.mkdir(parents=True, exist_ok=True)
OUT_DIR = Path("output")
OUT_DIR.mkdir(parents=True, exist_ok=True)
SUPPLEMENTS_SUBDIR_NAME = "supplements" # Subdirectory for supplements

# Europe PMC imposes a 30 req/min soft limit → be polite.
HEADERS_EPMC   = {"User-Agent": "PrimerMiner/0.3 (tmolley@ucsd.edu)"} # Updated version
BASE_URL_EPMC  = "https://www.ebi.ac.uk/europepmc/webservices/rest/"



# -----------------------------------------------------------------------------------
# 1.  Utility Functions (Europe PMC Search, HGNC, Date Ranges) ─────────────────────
# -----------------------------------------------------------------------------------
def get_weekly_date_ranges(year: int, month: int) -> List[Dict[str, any]]:
    """
    Generates a list of dictionaries, each representing a weekly period for querying.
    Each dictionary contains 'week_num', 'start_date', and 'end_date'.
    Weeks are defined to cover the entire month, with the last week potentially being shorter.
    """
    ranges = []
    month_start_date = datetime.date(year, month, 1)
    _, num_days_in_month = calendar.monthrange(year, month)
    month_end_date = datetime.date(year, month, num_days_in_month)

    current_query_start_date = month_start_date
    week_counter = 1
    while current_query_start_date <= month_end_date:
        # End of the query week is 6 days from start, capped at month end
        current_query_end_date = current_query_start_date + datetime.timedelta(days=6)
        if current_query_end_date > month_end_date:
            current_query_end_date = month_end_date
        
        ranges.append({
            "week_num": week_counter,
            "start_date": current_query_start_date,
            "end_date": current_query_end_date
        })
        
        # Next query week starts the day after the current one ends
        current_query_start_date = current_query_end_date + datetime.timedelta(days=1)
        week_counter += 1
    return ranges

# (Remove or comment out your old cleanup_downloaded_files function)

def remove_directory_recursively(dir_path: Path):
    """
    Deletes the specified directory and all its contents.
    """
    if dir_path.exists() and dir_path.is_dir():
        print(f"🧹 Deleting directory and its contents: {dir_path}")
        try:
            shutil.rmtree(dir_path)
            print(f"  ✅ Successfully deleted {dir_path}")
        except Exception as e:
            print(f"  ❌ Error deleting directory {dir_path}: {e}")
    # else:
    #     print(f"ℹ️  Directory not found for deletion (or not a directory): {dir_path}") # Optional: for debugging

def europepmc_search_all(query: str,
                         base_url: str,
                         headers: dict,
                         max_records: int = 10000000,
                         page_size: int = 1000,
                         throttle: float = 0.3) -> List[Dict]:
    hits, retrieved = [], 0
    cursor = "*"
    while retrieved < max_records:
        url = (f"{base_url}search"
               f"?query={requests.utils.quote(query)}"
               f"&format=json"
               f"&pageSize={page_size}"
               f"&cursorMark={cursor}")
        try:
            r = requests.get(url, headers=headers, timeout=40)
            r.raise_for_status()
            data = r.json()
        except requests.exceptions.RequestException as e:
            print(f"Europe PMC search request failed: {e}. Retrying once...")
            time.sleep(5) # Wait before retrying
            try:
                r = requests.get(url, headers=headers, timeout=60)
                r.raise_for_status()
                data = r.json()
            except requests.exceptions.RequestException as e_retry:
                raise RuntimeError(f"Europe PMC API error after retry: {e_retry}")


        if "resultList" not in data:
            raise RuntimeError(f"Europe PMC API error: {data.get('error', data)}")

        batch = data["resultList"]["result"]
        if not batch:
            break
        hits.extend(batch)
        retrieved += len(batch)
        cursor = data.get("nextCursorMark")
        if not cursor:
            break
        time.sleep(throttle)
    return hits[:max_records]

def load_hgnc_symbol_set(cache_dir_name="hgnc_cache") -> set[str]:
    cache_path = Path(cache_dir_name) / "homo_sapiens.gene_info.gz"
    if not cache_path.exists():
        cache_path.parent.mkdir(parents=True, exist_ok=True)
        url = ("https://ftp.ncbi.nlm.nih.gov/gene/DATA/GENE_INFO/"
               "Mammalia/Homo_sapiens.gene_info.gz")
        print("⏬  Downloading HGNC symbol list …")
        try:
            r = requests.get(url, timeout=60)
            r.raise_for_status()
            cache_path.write_bytes(r.content)
        except requests.exceptions.RequestException as e:
            print(f"Error downloading HGNC symbols: {e}. Proceeding without HGNC symbol validation.")
            return set()
    symbols = set()
    try:
        with gzip.open(cache_path, "rt", encoding='utf-8') as fh:
            for line in fh:
                if line.startswith("#"): continue
                cols = line.rstrip("\n").split("\t")
                if len(cols) > 2: symbols.add(cols[2].upper())
    except Exception as e:
        print(f"Error reading HGNC file {cache_path}: {e}. Proceeding without HGNC symbol validation.")
        return set()
    if not symbols: print(f"Warning: HGNC symbol set is empty from {cache_path}.")
    else: print(f"Loaded {len(symbols)} HGNC symbols.")
    return symbols

# ---------------------------------------------------------------------------
# 2.  File Handling and Downloading (XML, Supplements)
# ---------------------------------------------------------------------------
def get_or_download_xml(pmcid: str, dest_dir: Path, base_url: str, headers: dict) -> Path:
    fpath = dest_dir / f"{pmcid}.xml"
    if fpath.exists(): return fpath
    url = f"{base_url}{pmcid}/fullTextXML"
    r   = requests.get(url, headers=headers, timeout=30)
    r.raise_for_status()
    fpath.parent.mkdir(parents=True, exist_ok=True) # Ensure dest_dir exists
    fpath.write_text(r.text, encoding="utf-8")
    return fpath

def rename_supp_files(pmcid: str, file_paths: list[Path]) -> list[Path]:
    new_paths = []
    for i, old_path in enumerate(sorted(file_paths), 1):
        new_name = f"{pmcid}_supp{i}{old_path.suffix.lower()}"
        new_path = old_path.with_name(new_name)
        if new_path.exists() and old_path != new_path :
            timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S%f")
            new_name = f"{pmcid}_supp{i}_{timestamp}{old_path.suffix.lower()}"
            new_path = old_path.with_name(new_name)
            if new_path.exists():
                 new_paths.append(old_path)
                 continue
        try:
            if old_path != new_path: old_path.rename(new_path)
            new_paths.append(new_path)
        except OSError:
            if old_path.exists(): new_paths.append(old_path)
    return new_paths

def fetch_supplements_from_xml(pmcid: str,
                               xml_path: Path,
                               dest_dir: Path,
                               headers: dict,
                               pause: float = 0.2) -> list[Path]:
    dest_dir.mkdir(parents=True, exist_ok=True)
    if not xml_path.exists(): return []
    try:
        tree = etree.parse(str(xml_path))
        root = tree.getroot()
    except etree.XMLSyntaxError: # Handle cases where XML might be malformed
        # print(f"Warning: Could not parse XML for {pmcid} at {xml_path} to fetch supplements.")
        return []

    ns    = {"x": "http://www.w3.org/1999/xlink"}
    hrefs = set()
    hrefs.update(root.xpath("//supplementary-material/@x:href", namespaces=ns))
    hrefs.update(root.xpath("//supplementary-material//media/@x:href", namespaces=ns))
    hrefs.update(root.xpath("//ext-link[contains(@ext-link-type,'supplement')]/@x:href", namespaces=ns))
    if not hrefs: return []

    base_supp_url  = f"https://europepmc.org/articles/{pmcid}/bin/"
    saved = []
    for href in hrefs:
        url   = href if href.startswith("http") else base_supp_url + href.lstrip("/")
        fname = url.split("/")[-1].split("?")[0]
        if not fname.lower().endswith((".pdf", ".docx")): continue
        fpath = dest_dir / fname
        if fpath.exists():
            saved.append(fpath)
            continue
        try:
            r = requests.get(url, headers=headers, timeout=60)
            r.raise_for_status()
            fpath.write_bytes(r.content)
            saved.append(fpath)
            time.sleep(pause)
        except requests.HTTPError as e:
            print(f"⚠️  {pmcid} supp {fname} DL err: {e}")
        except Exception as e_gen:
            print(f"⚠️  {pmcid} supp {fname} gen err: {e_gen}")
    return saved

# -----------------------------------------------------------------------------------
# 3.  Content Parsing and Extraction Logic (Core primer finding)
# -----------------------------------------------------------------------------------
def get_element_text_content(element) -> str:
    return " ".join(text.strip() for text in element.xpath(".//text()") if text.strip()).strip()

def parse_xml_table_to_dataframe(table_etree_element):
    headers = []
    header_elements = table_etree_element.xpath("./thead/tr/th | ./tr[1]/th | ./thead/tr/td | ./tr[1]/td")
    if header_elements and all(el.tag == 'th' for el in table_etree_element.xpath("./tr[1]/*")):
        headers = [get_element_text_content(th).strip() for th in header_elements]
    elif header_elements:
        headers = [get_element_text_content(th).strip() for th in header_elements]
    rows_data = []
    tr_elements = table_etree_element.xpath("./tbody/tr | ./tr")
    start_row_index = 0
    if headers and len(tr_elements) > 0:
        first_tr_cells = tr_elements[0].xpath("./th | ./td")
        if [get_element_text_content(c).strip() for c in first_tr_cells] == headers:
            start_row_index = 1
    for r_idx in range(start_row_index, len(tr_elements)):
        cells = tr_elements[r_idx].xpath("./td | ./th")
        row_text = [get_element_text_content(cell).strip() for cell in cells]
        if any(rt.strip() for rt in row_text): rows_data.append(row_text)
    if not rows_data: return pd.DataFrame()
    df = pd.DataFrame(rows_data)
    if headers and len(headers) == df.shape[1]: df.columns = headers
    elif not df.empty and not headers and df.shape[0] > 1 and \
         not any(is_valid_primer_sequence(extract_sequence_and_direction_from_cell(str(c))[0] or "",18,40) for c in df.iloc[0]):
        df.columns = [str(c).strip() for c in df.iloc[0]]
        df = df[1:].reset_index(drop=True)
    return df

def scan_dataframe_for_primers(df: pd.DataFrame, pmcid: str,
                               source_description: str,
                               source_filename_context: str,
                               hgnc_symbols: set) -> List[Dict]:
    found_primers_in_df = []
    if df.empty: return found_primers_in_df
    df_headers_original = [str(col).strip() for col in df.columns]
    df_headers_lower = [h.lower() for h in df_headers_original]

    for r_idx in range(df.shape[0]):
        row_context_gene = "Unknown"
        if df.shape[1] > 0:
            first_cell = str(df.iloc[r_idx, 0]).strip()
            if first_cell and first_cell.lower() != 'nan':
                bases, _, _ = extract_sequence_and_direction_from_cell(first_cell)
                if not (bases and is_valid_primer_sequence(bases, min_len=15)):
                    name = clean_gene_name(first_cell)
                    if name != "Unknown" and (name.upper() in hgnc_symbols or
                       (len(name) > 4 and not any(k in name.lower() for k in FORWARD_KEYWORDS + REVERSE_KEYWORDS))):
                        row_context_gene = name
        for c_idx in range(df.shape[1]):
            try:
                cell_text = str(df.iloc[r_idx, c_idx]).strip()
                if not cell_text or cell_text.lower() == 'nan' or len(cell_text) < 10: continue
                bases, direction, original_text = extract_sequence_and_direction_from_cell(cell_text)
                if bases and is_valid_primer_sequence(bases):
                    seq, p_dir, p_gene = bases, direction or "Unknown", row_context_gene
                    if p_gene == "Unknown" and c_idx > 0:
                        left_cell_text = str(df.iloc[r_idx, c_idx - 1]).strip()
                        if left_cell_text and left_cell_text.lower() != 'nan':
                            lbases, _, _ = extract_sequence_and_direction_from_cell(left_cell_text)
                            if not (lbases and is_valid_primer_sequence(lbases, min_len=15)):
                                lname = clean_gene_name(left_cell_text)
                                if lname != "Unknown": p_gene = lname
                    if p_gene == "Unknown" or len(p_gene) < 2: # Simplified check
                        match = re.match(r"([a-zA-Z0-9_.\-]+(?:[-/][a-zA-Z0-9_.\-]+)*)\s*[:\-(]", original_text)
                        if match:
                            prefix_name = clean_gene_name(match.group(1))
                            if prefix_name != "Unknown": p_gene = prefix_name
                    if p_dir == "Unknown" and c_idx < len(df_headers_lower):
                        header = df_headers_lower[c_idx]
                        if any(k in header for k in FORWARD_KEYWORDS): p_dir = "Forward"
                        elif any(k in header for k in REVERSE_KEYWORDS): p_dir = "Reverse"
                    if p_gene == "Unknown" and c_idx < len(df_headers_lower):
                        header_orig = df_headers_original[c_idx]
                        header_low = df_headers_lower[c_idx]
                        if not any(k in header_low for k in FORWARD_KEYWORDS + REVERSE_KEYWORDS + ['sequence','primer','oligo','probe']) and len(header_orig) > 1:
                            h_gene = clean_gene_name(header_orig)
                            if h_gene != "Unknown": p_gene = h_gene
                    found_primers_in_df.append({
                        "PMCID": pmcid, "Gene": p_gene, "Sequence": seq,
                        "Orientation": p_dir, "Source File": source_filename_context,
                        "Page": source_description, "Original Cell Text": original_text})
            except Exception: continue
    return found_primers_in_df

def parse_docx_table_to_dataframe(docx_table_obj):
    data, keys = [], None
    if docx_table_obj.rows:
        cells = [cell.text.strip() for cell in docx_table_obj.rows[0].cells]
        if cells and sum(1 for h in cells if not h) < len(cells)/2 and all(len(h) < 100 for h in cells):
            keys = cells
    for i, row in enumerate(docx_table_obj.rows):
        text = [cell.text.strip() for cell in row.cells]
        if keys and i == 0 and len(docx_table_obj.rows) > 1: continue
        if keys and len(text) == len(keys) and (i > 0 or len(docx_table_obj.rows) == 1):
            data.append(dict(zip(keys, text)))
        else: data.append(text)
    if not data: return pd.DataFrame()
    try:
        df = pd.DataFrame(data)
        if keys and not isinstance(data[0], dict) and len(keys) == df.shape[1]: df.columns = keys
    except Exception: df = pd.DataFrame()
    if not df.empty and all(isinstance(c,int) for c in df.columns) and df.shape[0]>0:
        first_row = [str(c).strip() for c in df.iloc[0]]
        if df.shape[0]>1 and not any(is_valid_primer_sequence(extract_sequence_and_direction_from_cell(str(c))[0] or "",18,40) for c in first_row) and any(len(str(c))>0 for c in first_row):
            df.columns, df = first_row, df[1:].reset_index(drop=True)
    return df

def extract_primers_from_docx(docx_path: Path, pmcid: str, hgnc_symbols: set) -> List[Dict]:
    found_primers = []
    if not docx_path.exists(): return found_primers
    try: doc = docx.Document(docx_path)
    except Exception: return found_primers
    if doc.tables:
        for i, table in enumerate(doc.tables):
            df = parse_docx_table_to_dataframe(table)
            if not df.empty:
                primers = scan_dataframe_for_primers(df, pmcid, f"DOCX Table {i+1}", docx_path.name, hgnc_symbols)
                if primers: found_primers.extend(primers)
    unique_prose, dna_regex = set(), re.compile(r'\b((?:[ACGUTNRYKMSWBDHV]\s*){15,50})\b', re.IGNORECASE)
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or len(text) < 20 or text in unique_prose: continue
        unique_prose.add(text)
        for match in dna_regex.finditer(text):
            raw_dna = match.group(1)
            bases, direction, _ = extract_sequence_and_direction_from_cell(raw_dna)
            if bases and is_valid_primer_sequence(bases):
                seq, p_dir, p_gene = bases, direction or "Unknown", "Unknown"
                start, end = match.start(), match.end()
                ctx_chars = 80
                before = text[max(0, start - ctx_chars) : start]
                after = text[end : min(len(text), end + ctx_chars)]
                if p_dir == "Unknown":
                    if any(k in before.lower() or k in after.lower() for k in FORWARD_KEYWORDS): p_dir = "Forward"
                    elif any(k in before.lower() or k in after.lower() for k in REVERSE_KEYWORDS): p_dir = "Reverse"
                gene_ctx = before + raw_dna + after
                pot_genes = re.findall(r'\b([A-Z][A-Za-z0-9-]{2,15})\b', gene_ctx)
                best_dist, noise = float('inf'), {"PCR","DNA","RNA","PRIMER"} | {k.upper() for k in FORWARD_KEYWORDS+REVERSE_KEYWORDS}
                for pg_raw in pot_genes:
                    pg_up = pg_raw.upper()
                    if (pg_up in hgnc_symbols or (len(pg_raw)>2 and pg_up not in noise and not pg_raw.islower())):
                        try:
                            for m in re.finditer(re.escape(pg_raw), gene_ctx):
                                dist = abs(m.start() - len(before))
                                if dist < best_dist and dist < ctx_chars:
                                    best_dist, name = dist, clean_gene_name(pg_raw)
                                    if name != "Unknown": p_gene = name
                        except: pass
                context_snip = f"...{text[max(0,start-30):min(len(text),end+30)]}..."
                found_primers.append({"PMCID":pmcid,"Gene":p_gene,"Sequence":seq,"Orientation":p_dir,
                                      "Source File":docx_path.name,"Page":"DOCX Paragraph","Original Cell Text":context_snip})
    return found_primers

def extract_pmcid_from_filename(filename: str) -> str:
    if not isinstance(filename, str): return "UnknownPMCID"
    m = re.match(r"^(.*?)_supp(?:lemental|lement|s)?\d*\.(pdf|docx)$", filename, re.IGNORECASE)
    if m: return m.group(1).upper().startswith("PMC") and m.group(1) or m.group(1)
    m = re.match(r"^(PMC\d+)\.(pdf|docx)$", filename, re.IGNORECASE)
    if m: return m.group(1)
    return os.path.splitext(filename)[0]

def extract_sequence_and_direction_from_cell(cell_text: str) -> Tuple[Optional[str], Optional[str], str]:
    if not isinstance(cell_text, str) or not cell_text.strip(): return None, None, ""
    text, orig_text, direction = cell_text.strip(), cell_text.strip(), None
    affixes = {"Forward": [(r"^(?:FP[\s:-]*|FWD[\s:-]*|FORWARD[\s:-]*|SENSE[\s:-]*)", "prefix"), (r"\s*\((?:Forward|FWD|FP|Sense)\)$", "suffix")],
               "Reverse": [(r"^(?:RP[\s:-]*|REV[\s:-]*|REVERSE[\s:-]*|ANTISENSE[\s:-]*)", "prefix"), (r"\s*\((?:Reverse|REV|RP|Antisense)\)$", "suffix")]}
    p5, p3 = r"^(?:5['`´‘’]?[-–—]?)", r"(?:[-–—]?3['`´‘’]?)$"
    changed = True
    while changed:
        text_before, changed, dir_found = text, False, False
        for key, patterns in affixes.items():
            if dir_found and direction: break
            for pattern, affix_type in patterns:
                len_before = len(text)
                if affix_type == "prefix" and (m:=re.match(pattern,text,re.IGNORECASE)):
                    if not direction: direction=key
                    text, dir_found = text[m.end():].strip(), True
                elif affix_type == "suffix" and (m:=re.search(pattern,text,re.IGNORECASE)):
                    if not (m.start()==0 and direction):
                        if not direction: direction=key
                        text = text[:m.start()].strip()
                if len(text) != len_before:
                    changed = True
                    if dir_found: break
            if dir_found and direction: break
    text = re.sub(p5,"",text,flags=re.IGNORECASE).strip()
    text = re.sub(p3,"",text,flags=re.IGNORECASE).strip()
    candidates = re.findall(r'([ACTG]{10,})', text, re.IGNORECASE)
    if not candidates: return None, direction, orig_text
    best_seq = max((c for c in candidates if is_valid_primer_sequence(c)), key=len, default="")
    return best_seq.upper() if best_seq else None, direction, orig_text

def is_valid_primer_sequence(sequence_bases: str, min_len=18, max_len=40) -> bool:
    return isinstance(sequence_bases,str) and min_len <= len(sequence_bases) <= max_len and \
           all(c.upper() in ALLOWED_DNA_CHARS_FOR_FILTER for c in sequence_bases)

def clean_gene_name(gene_text: str) -> str:
    if not isinstance(gene_text,str) or not gene_text.strip() or gene_text.lower()=='nan': return "Unknown"
    m = re.match(r"([a-zA-Z0-9_.\-]+(?:[-/][a-zA-Z0-9_.\-]+)*)", gene_text.strip())
    if m:
        name = m.group(1)
        if len(name)>1 and not (name.isnumeric() and len(name)>4) and \
           name.lower() not in (FORWARD_KEYWORDS+REVERSE_KEYWORDS+["sense","antisense","fp","rp"]) and \
           not (len(name)>15 and all(c.upper() in ALLOWED_DNA_CHARS_FOR_FILTER for c in name)):
            return name
    return "Unknown"

def extract_primers_from_xml(xml_path: Path, pmcid: str, hgnc_symbols: set) -> List[Dict]:
    found_primers = []
    if not xml_path.exists(): return found_primers
    try:
        with open(xml_path, 'rb') as f: content = f.read()
        root = etree.fromstring(content, etree.XMLParser(recover=True,strip_cdata=True,resolve_entities=False,no_network=True))
    except Exception: return found_primers # Error parsing XML
    tables = root.xpath("//table[ancestor::table-wrap]") or root.xpath("//table")
    for i, table_el in enumerate(tables):
        df = parse_xml_table_to_dataframe(table_el)
        if not df.empty:
            primers = scan_dataframe_for_primers(df, pmcid, f"XML Table {i+1}", xml_path.name, hgnc_symbols)
            if primers: found_primers.extend(primers)
    prose_paths = ["//sec[translate(@sec-type,'METHODS','methods')]//p[not(.//table)]",
                   "//sec[translate(@sec-type,'MATERIALSMETHODS','materialsmethods')]//p[not(.//table)]",
                   "//body//p[not(ancestor::table-wrap) and not(ancestor::fig) and not(.//table)]",
                   "//fig//caption//p[not(.//table)]", "//table-wrap//caption//p[not(.//table)]",
                   "//table-wrap-foot//p[not(.//table)]"]
    unique_prose, dna_regex = set(), re.compile(r'\b((?:[ACGUTNRYKMSWBDHV]\s*){15,50})\b', re.IGNORECASE)
    for xp in prose_paths:
        try:
            for elem in root.xpath(xp):
                text = get_element_text_content(elem)
                if not text or len(text) < 20 or text in unique_prose: continue
                unique_prose.add(text)
                for match in dna_regex.finditer(text):
                    raw_dna = match.group(1)
                    bases, direction, _ = extract_sequence_and_direction_from_cell(raw_dna)
                    if bases and is_valid_primer_sequence(bases):
                        seq, p_dir, p_gene = bases, direction or "Unknown", "Unknown"
                        s, e = match.start(), match.end()
                        ctx = 80
                        before, after = text[max(0,s-ctx):s], text[e:min(len(text),e+ctx)]
                        if p_dir == "Unknown":
                            if any(k in before.lower() or k in after.lower() for k in FORWARD_KEYWORDS): p_dir = "Forward"
                            elif any(k in before.lower() or k in after.lower() for k in REVERSE_KEYWORDS): p_dir = "Reverse"
                        gene_ctx = before + raw_dna + after
                        pot_genes = re.findall(r'\b([A-Z][A-Za-z0-9-]{2,15})\b', gene_ctx)
                        best_d, noise = float('inf'), {"PCR","DNA","RNA","PRIMER"} | {k.upper() for k in FORWARD_KEYWORDS+REVERSE_KEYWORDS}
                        for pg_raw in pot_genes:
                            pg_up = pg_raw.upper()
                            if pg_up in hgnc_symbols or (len(pg_raw)>2 and pg_up not in noise and not pg_raw.islower()):
                                for m_pg in re.finditer(re.escape(pg_raw),gene_ctx,re.IGNORECASE):
                                    d = abs(m_pg.start() - len(before))
                                    if d < best_d and d < ctx:
                                        best_d, name = d, clean_gene_name(pg_raw)
                                        if name != "Unknown": p_gene = name
                        snip = f"...{text[max(0,s-30):min(len(text),e+30)]}..."
                        found_primers.append({"PMCID":pmcid,"Gene":p_gene,"Sequence":seq,"Orientation":p_dir,
                                             "Source File":xml_path.name,"Page":"Main Text Prose","Original Cell Text":snip})
        except Exception: continue
    return found_primers

def find_primers_sequence_first(pdf_path_str: str, page_spec="all", hgnc_symbols: set = set()) -> List[Dict]:
    all_primers = []
    pdf_name = os.path.basename(pdf_path_str)
    camelot_timeout = 30 # Timeout for each Camelot call

    # Parameters for Camelot
    # Increased line_scale for lattice, as default sometimes misses lines in dense tables
    lattice_params = {'pages': page_spec, 'flavor': 'lattice', 'line_scale': 40, 
                      'shift_text': [' '], 'copy_text': ['v'], 'suppress_stdout': True}
    stream_params = {'pages': page_spec, 'flavor': 'stream', 'suppress_stdout': True}
    
    extracted_tables_camelot = []

    # Run Camelot calls sequentially within a ProcessPool to isolate them and manage timeouts
    # Using max_workers=1 ensures Camelot runs one at a time, which can be more stable.
    # Pebble ProcessPool is used here as per your imports for robust timeouts.
    with ProcessPool(max_workers=1, context=get_context('spawn') if os.name != 'nt' else None) as pool: # Use spawn context if not windows for safety
        try:
            # print(f"    Trying Camelot (Lattice) for {pdf_name}, pages: {page_spec}")
            future_lattice = pool.schedule(camelot.read_pdf, args=[pdf_path_str], kwargs=lattice_params, timeout=camelot_timeout)
            lattice_table_list = future_lattice.result() # Wait for result
            if lattice_table_list and lattice_table_list.n > 0:
                extracted_tables_camelot.extend(list(lattice_table_list))
        except PebbleTimeoutError:
            print(f"    Lattice TIMEOUT for {pdf_name} (pages '{page_spec}') after {camelot_timeout}s")
        except Exception as e:
            # Catch broad exceptions from Camelot, including ghostscript errors if it's called by camelot
            print(f"    Lattice FAIL for {pdf_name} (pages '{page_spec}'): {type(e).__name__} {str(e)[:100]}")

        try:
            # print(f"    Trying Camelot (Stream) for {pdf_name}, pages: {page_spec}")
            future_stream = pool.schedule(camelot.read_pdf, args=[pdf_path_str], kwargs=stream_params, timeout=camelot_timeout)
            stream_table_list = future_stream.result() # Wait for result
            if stream_table_list and stream_table_list.n > 0:
                extracted_tables_camelot.extend(list(stream_table_list))
        except PebbleTimeoutError:
            print(f"    Stream TIMEOUT for {pdf_name} (pages '{page_spec}') after {camelot_timeout}s")
        except Exception as e:
            print(f"    Stream FAIL for {pdf_name} (pages '{page_spec}'): {type(e).__name__} {str(e)[:100]}")

    if not extracted_tables_camelot:
        # print(f"    No tables extracted by Camelot for {pdf_name}, pages {page_spec}.")
        return []

    unique_table_hashes = set() # To avoid processing identical tables from different methods
    pmcid_from_file = extract_pmcid_from_filename(pdf_name)

    for i, table_obj in enumerate(extracted_tables_camelot):
        try:
            # Create a hashable representation of the table to detect duplicates
            table_hash = (table_obj.page, table_obj.df.shape, table_obj.df.iloc[0,0] if not table_obj.df.empty else None, table_obj.flavor)
            if table_hash in unique_table_hashes:
                continue
            unique_table_hashes.add(table_hash)
        except Exception: pass # Ignore hashing errors, just process the table

        df_from_camelot = table_obj.df
        if df_from_camelot.empty:
            continue
            
        primers_from_df = scan_dataframe_for_primers(
            df_from_camelot, 
            pmcid_from_file, 
            f"PDF Pg {table_obj.page}, Tbl {i+1} ({table_obj.flavor})", 
            pdf_name, 
            hgnc_symbols
        )
        if primers_from_df:
            all_primers.extend(primers_from_df)
            
    return all_primers

# -----------------------------------------------------------------------------------
# 4.  Single Paper Processing Function (for parallel execution) ─────────────────────
# -----------------------------------------------------------------------------------
# Custom exception for internal timeout
class InternalWorkerTimeoutError(Exception):
    pass

def _internal_timeout_handler(signum, frame):
    """Signal handler that raises our custom timeout exception."""
    raise InternalWorkerTimeoutError("Internal paper processing time limit exceeded.")

def process_single_paper(
    hit_info: Dict,
    processing_data_dir: Path,
    hgnc_symbols_set: set,
    n_last_pages_supp: int,
    throttle_sec: float,
    base_url_epmc_global: str,
    headers_epmc_global: dict,
    internal_timeout_duration: int = 28 # Internal timeout in seconds (e.g., slightly less than main timeout)
) -> Tuple[List[Dict], int, Optional[str], float]:

    overall_processing_start_time = time.time() # For accurate duration tracking

    # --- Signal-based Timeout Setup (Unix-like systems) ---
    original_sigalrm_handler = None
    can_use_signals = hasattr(signal, 'SIGALRM')

    if can_use_signals:
        try:
            original_sigalrm_handler = signal.getsignal(signal.SIGALRM)
            signal.signal(signal.SIGALRM, _internal_timeout_handler)
            signal.alarm(internal_timeout_duration)
        except ValueError: #  Happens if not in main thread on some OS, or other issues
            can_use_signals = False # Fallback to no signal-based timeout
            if original_sigalrm_handler is not None: # Restore if partially set
                 signal.signal(signal.SIGALRM, original_sigalrm_handler)
            print(f"    ⚠️ Could not set SIGALRM handler for {hit_info.get('pmcid')}, signal-based timeout disabled for this task.")


    primers_for_paper = []
    count = 0
    pmcid = hit_info.get("pmcid") or hit_info.get("id")
    # Default to an error PMCID if none is found, for logging purposes.
    pmcid_for_log = pmcid if pmcid else f"UNKNOWN_ID_IN_HIT_{hit_info.get('id', 'NO_ID')}"


    try:
        # --- Early exit for invalid PMCID ---
        if not (pmcid and pmcid.startswith("PMC")):
            duration = time.time() - overall_processing_start_time
            return [], 0, pmcid_for_log, duration

        # --- Directory Setup ---
        supp_dir = processing_data_dir / SUPPLEMENTS_SUBDIR_NAME
        processing_data_dir.mkdir(parents=True, exist_ok=True)
        supp_dir.mkdir(parents=True, exist_ok=True)
        citations = hit_info.get("citedByCount", 0)

        # --- Main Processing Logic ---
        # 1. XML Processing
        # Pass a shorter timeout to network calls within helper functions
        # (e.g. internal_timeout_duration - 5 or a fixed value like 15s)
        xml_file_timeout = max(5, internal_timeout_duration - 10) # Example: ensure at least 5s, up to T-10s
        xml_file = get_or_download_xml(pmcid, processing_data_dir, base_url_epmc_global, headers_epmc_global)
        time.sleep(throttle_sec) # Brief pause

        if xml_file.exists():
            xml_primers = extract_primers_from_xml(xml_file, pmcid, hgnc_symbols_set)
            if xml_primers:
                for entry in xml_primers: entry["Citation Count"] = citations
                primers_for_paper.extend(xml_primers)
                count += len(xml_primers)

        # 2. Supplements Fetching (conditional)
        supp_to_process = []
        # Check elapsed time cooperatively before potentially slow supplement fetching
        if time.time() - overall_processing_start_time > internal_timeout_duration -5 : # -5 to leave buffer
             raise InternalWorkerTimeoutError("Cooperative timeout before fetching supplements.")

        if count <= 8 and xml_file.exists():
            
            dl_supps = fetch_supplements_from_xml(pmcid, xml_file, supp_dir, headers_epmc_global, pause=0.2)
            if dl_supps:
                supp_to_process = rename_supp_files(pmcid, dl_supps)
        
        # 3. Supplements Processing Loop
        for supp_idx, supp_path in enumerate(supp_to_process):
            if count > 8: break

            # Cooperative check before processing each supplement
            if time.time() - overall_processing_start_time > internal_timeout_duration:
                raise InternalWorkerTimeoutError(f"Cooperative timeout processing supplement {supp_idx+1} ({supp_path.name}).")

            supp_primers_current, page_spec, process_this_supp = [], "all", True
            s_path_str, s_suffix = str(supp_path), supp_path.suffix.lower()

            if s_suffix == ".pdf":
                try:
                    # pdfplumber operations can be slow. SIGALRM is the main defense here.
                    # For very problematic PDFs, pdfplumber itself might hang in C code.
                    with pdfplumber.open(supp_path) as pdf:
                        if not pdf.pages:
                            process_this_supp = False
                        else:
                            total_pg = len(pdf.pages)
                            if total_pg > 0:
                                # This extract_text can be very slow for complex pages/PDFs
                                first_pg_txt = pdf.pages[0].extract_text(x_tolerance=1, y_tolerance=1) or ""
                                if "reporting summary" in first_pg_txt.lower(): page_spec = "1"
                                elif total_pg <= n_last_pages_supp: page_spec = f"1-{total_pg}"
                                else: page_spec = f"{max(1, total_pg - n_last_pages_supp + 1)}-{total_pg}"
                            else:
                                process_this_supp = False
                except InternalWorkerTimeoutError: # Re-raise if SIGALRM triggered during PDF
                    raise
                except Exception as pdf_err: # Catch other PDF processing errors
                    print(f"    Error opening/reading PDF {supp_path.name} for {pmcid}: {pdf_err}")
                    page_spec = "all" # Original behavior, try to process all pages if opening failed this way
                    # process_this_supp remains True to attempt find_primers_sequence_first
                
                if process_this_supp:
                    supp_primers_current = find_primers_sequence_first(s_path_str, page_spec, hgnc_symbols_set)

            elif s_suffix == ".docx": # 'process_this_supp' is not used for docx in original logic
                supp_primers_current = extract_primers_from_docx(supp_path, pmcid, hgnc_symbols_set)
            
            if supp_primers_current:
                for entry in supp_primers_current: entry["Citation Count"] = citations
                primers_for_paper.extend(supp_primers_current)
                count += len(supp_primers_current)
        
        # If we've reached here, processing completed normally (or partially if count > 8)
        # within the time limit.

    except requests.HTTPError as e:
        # Handle specific, known operational errors like HTTP errors
        print(f"    HTTPError for {pmcid_for_log}: {e}")
        # primers_for_paper, count remain as they are.
    except InternalWorkerTimeoutError as e_timeout:
        # This catches timeouts raised by SIGALRM or cooperative checks
        print(f"    ⏳ INTERNAL TIMEOUT for {pmcid_for_log} after ~{internal_timeout_duration}s. Details: {e_timeout}")
        # primers_for_paper, count will contain whatever was gathered before timeout.
    except Exception as e:
        # Catch any other unexpected errors during processing
        print(f"    Error processing {pmcid_for_log}: {type(e).__name__} - {e}")
        # primers_for_paper, count remain as they are.
    
    finally:
        # --- Crucial: Cleanup signal handler and alarm ---
        if can_use_signals and original_sigalrm_handler is not None:
            signal.alarm(0) # Disable any pending alarm
            signal.signal(signal.SIGALRM, original_sigalrm_handler) # Restore original

    duration = time.time() - overall_processing_start_time
    return primers_for_paper, count, pmcid_for_log, duration

# -----------------------------------------------------------------------------------
# 5. Main Execution Block ───────────────────────────────────────────────────────────
# -----------------------------------------------------------------------------------
if __name__ == "__main__":
    SCRIPT_ROOT = Path(__file__).resolve().parent
    TMP_DIR_FOR_SCRIPT = SCRIPT_ROOT / "script_temp_files"
    try:
        TMP_DIR_FOR_SCRIPT.mkdir(parents=True, exist_ok=True)
        os.environ['TMPDIR'] = str(TMP_DIR_FOR_SCRIPT)
        os.environ['TEMP'] = str(TMP_DIR_FOR_SCRIPT)
        os.environ['TMP'] = str(TMP_DIR_FOR_SCRIPT)
        print(f"ℹ️  Attempting to use {TMP_DIR_FOR_SCRIPT} for temporary files.")
    except Exception as e:
        print(f"⚠️ Could not create or set script_temp_files directory: {e}")

    DATA_DIR_BASE = Path("data/") # Changed for testing
    OUT_DIR = Path("output/")       # Changed for testing
    SUPPLEMENTS_SUBDIR_NAME = "supplements"
    DATA_DIR_BASE.mkdir(parents=True, exist_ok=True)
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    MAX_RECORDS_PER_QUERY = 1000000
    N_LAST_PAGES_SUPPLEMENT = 5
    NUM_PARALLEL_WORKERS = 32
    MAIN_XML_DL_THROTTLE = 0.2
    INTERNAL_WORKER_TIMEOUT_SECONDS = 28 
    SINGLE_PAPER_PROCESSING_TIMEOUT = 28
    CFG_PROCESS_START_YEAR = 2000
    CFG_PROCESS_END_YEAR = 1995 # Reduced for quicker test run
    CFG_RESUME_MONTH = 1
    CFG_RESUME_WEEK = 1

    print(f"🧬 Configuration: Processing from Year {CFG_PROCESS_START_YEAR}, Month {CFG_RESUME_MONTH}, Week {CFG_RESUME_WEEK}")
    print(f"🧬 Down to Year {CFG_PROCESS_END_YEAR} (inclusive).")
    print("Loading HGNC symbols...")
    HGNC_SYMBOLS = load_hgnc_symbol_set()

    overall_start_time = time.time()
    total_papers_overall = 0
    total_primers_overall = 0
    total_papers_timed_out_overall = 0
    total_papers_other_exceptions_overall = 0 # New counter

    for year in range(CFG_PROCESS_START_YEAR, CFG_PROCESS_END_YEAR - 1, -1):
        # ... (year/month/week skipping logic remains the same) ...
        effective_start_month_for_this_year = 1
        if year == CFG_PROCESS_START_YEAR:
            effective_start_month_for_this_year = CFG_RESUME_MONTH

        for month in range(1, 13):
            if year == CFG_PROCESS_START_YEAR and month < effective_start_month_for_this_year:
                # print(f"⏭️  SKIPPING {year}-{month:02d} (before configured resume month {effective_start_month_for_this_year})")
                continue

            weekly_ranges = get_weekly_date_ranges(year, month)
            for week_info in weekly_ranges:
                week_num = week_info['week_num']
                start_date_week = week_info['start_date']
                end_date_week = week_info['end_date']

                if year == CFG_PROCESS_START_YEAR and \
                   month == effective_start_month_for_this_year and \
                   week_num < CFG_RESUME_WEEK:
                    # print(f"⏭️  SKIPPING {year}-{month:02d}-Week{week_num} (before configured resume week {CFG_RESUME_WEEK})")
                    continue

                batch_start_time = time.time()
                print(f"\n{'='*70}\n🚀 STARTING: {year}-{month:02d}-Week{week_num} ({start_date_week.isoformat()} to {end_date_week.isoformat()})\n{'='*70}\n")

                current_week_data_dir = DATA_DIR_BASE / str(year) / f"{month:02d}" / f"week_{week_num}"
                remove_directory_recursively(current_week_data_dir)
                current_week_data_dir.mkdir(parents=True, exist_ok=True)

                master_primers_week = []
                base_query = 'HAS_FT:y AND (METHODS:"qPCR" OR "RT-PCR" OR "real time PCR" OR "quantitative PCR")'
                query = f"{base_query} AND (FIRST_PDATE:[{start_date_week.isoformat()} TO {end_date_week.isoformat()}])"

                print(f"Searching Europe PMC: {query}")
                epmc_hits_week = []
                try:
                    epmc_hits_week = europepmc_search_all(query, BASE_URL_EPMC, HEADERS_EPMC, MAX_RECORDS_PER_QUERY, throttle=0.3)
                    print(f"Retrieved {len(epmc_hits_week)} records for this week.")
                except Exception as e:
                    print(f"Search failed for {year}-{month:02d}-Week{week_num}: {e}")
                    continue

                if not epmc_hits_week:
                    print(f"No records found for {year}-{month:02d}-Week{week_num}.")
                    continue

                (current_week_data_dir / SUPPLEMENTS_SUBDIR_NAME).mkdir(parents=True, exist_ok=True)

                papers_this_week_processed_successfully = 0
                papers_this_week_timed_out = 0
                papers_this_week_other_exceptions = 0 # New counter
                task_time_week_successful = 0
                
                future_to_pmcid_map = {}
                tqdm.write(f"ℹ️ Week {year}-{month:02d}-W{week_num}: Entering ProcessPoolExecutor context...")
                with ProcessPoolExecutor(max_workers=NUM_PARALLEL_WORKERS) as executor:
                    tqdm.write(f"ℹ️ Week {year}-{month:02d}-W{week_num}: ProcessPoolExecutor started. Submitting tasks...")
                    future_to_pmcid_map = {} # Ensure this is defined before use
                    for hit_idx, hit in enumerate(epmc_hits_week):
                        pmcid_for_submission = hit.get("pmcid") or hit.get("id") or f"UNKNOWN_PMCID_HIT_{hit_idx}"
                        submitted_future = executor.submit(process_single_paper, hit, current_week_data_dir,
                                                           HGNC_SYMBOLS, N_LAST_PAGES_SUPPLEMENT, MAIN_XML_DL_THROTTLE,
                                                           BASE_URL_EPMC, HEADERS_EPMC,internal_timeout_duration=INTERNAL_WORKER_TIMEOUT_SECONDS)
                        future_to_pmcid_map[submitted_future] = pmcid_for_submission
                    
                    tqdm.write(f"ℹ️ Week {year}-{month:02d}-W{week_num}: All {len(future_to_pmcid_map)} tasks submitted. Starting as_completed loop...")
                    total_tasks_for_week = len(future_to_pmcid_map)
                    tasks_fully_accounted_for_this_week = 0

                    failsafe_tasks_processed_target = total_tasks_for_week - 0
                    apply_failsafe_check = total_tasks_for_week > 5

                    # Keep track of all submitted futures to attempt cancellation if failsafe triggers
                    all_submitted_futures = list(future_to_pmcid_map.keys())

                    for future_idx, future in enumerate(tqdm(as_completed(all_submitted_futures), total=total_tasks_for_week, desc=f"Proc. {year}-{month:02d}-W{week_num}")):
                        pmid_for_log = future_to_pmcid_map[future]
                        
                        # --- FAILSAFE CHECK (at the beginning of processing each future) ---
                        if apply_failsafe_check and tasks_fully_accounted_for_this_week >= failsafe_tasks_processed_target:
                            tqdm.write(f"⚠️ FAILSAFE TRIGGERED for week {year}-{month:02d}-W{week_num}:")
                            tqdm.write(f"   Accounted for {tasks_fully_accounted_for_this_week}/{total_tasks_for_week} tasks.")
                            tqdm.write(f"   Target was to account for {failsafe_tasks_processed_target} tasks before giving up.")
                            tqdm.write(f"   Skipping remaining (up to {total_tasks_for_week - tasks_fully_accounted_for_this_week}) tasks.")
                            
                            # Attempt to cancel futures that haven't completed yet.
                            # This is best-effort and might not stop already running tasks.
                            cancelled_count = 0
                            # Iterate over the original list of all futures
                            for fut_to_cancel in all_submitted_futures:
                                if not fut_to_cancel.done(): # Check if it's not done already
                                    if fut_to_cancel.cancel(): # cancel() returns True if successfully cancelled
                                        cancelled_count += 1
                            if cancelled_count > 0:
                                tqdm.write(f"   Attempted to cancel {cancelled_count} pending/running tasks.")
                            else:
                                tqdm.write(f"   No tasks were pending cancellation or could be cancelled.")
                            
                            tqdm.write(f"ℹ️ Week {year}-{month:02d}-W{week_num}: Executing BREAK from as_completed loop due to failsafe.")
                            break # Exit the as_completed loop

                        # Process the current future
                        try:
                            # ... (your existing future.result() logic and success handling)
                            primers, count, pmid_from_result, duration = future.result(timeout=SINGLE_PAPER_PROCESSING_TIMEOUT)
                            if primers:
                                master_primers_week.extend(primers)
                            papers_this_week_processed_successfully += 1
                            task_time_week_successful += duration
                            if count > 0:
                                avg_success_time = task_time_week_successful / papers_this_week_processed_successfully if papers_this_week_processed_successfully > 0 else 0
                                tqdm.write(f"📄 OK: {pmid_from_result or pmid_for_log} ({duration:.2f}s). Found {count}. ({(future_idx + 1)}/{total_tasks_for_week}). Avg success: {avg_success_time:.2f}s")
                        
                        except TimeoutError: # From concurrent.futures
                            # ... (your existing TimeoutError handling)
                            tqdm.write(f"⏰ TIMEOUT for PMCID {pmid_for_log} after {SINGLE_PAPER_PROCESSING_TIMEOUT}s. ({(future_idx + 1)}/{total_tasks_for_week})")
                            papers_this_week_timed_out += 1
                            total_papers_timed_out_overall +=1
                        
                        except concurrent.futures.CancelledError:
                            # This will now be caught if future.cancel() above leads to it
                            tqdm.write(f"🚫 CANCELLED: PMCID {pmid_for_log} due to failsafe or other cancellation. ({(future_idx + 1)}/{total_tasks_for_week})")
                            papers_this_week_other_exceptions +=1 
                            total_papers_other_exceptions_overall +=1

                        except Exception as e_future:
                            # ... (your existing general Exception handling)
                            tqdm.write(f"❌ ERROR for PMCID {pmid_for_log} ({(future_idx + 1)}/{total_tasks_for_week}): {type(e_future).__name__} - {e_future}")
                            papers_this_week_other_exceptions +=1
                            total_papers_other_exceptions_overall +=1
                        
                        finally:
                            # This task has now been "accounted for"
                            tasks_fully_accounted_for_this_week += 1
                    
                    # This message prints when the as_completed loop finishes, either normally or via the break
                    tqdm.write(f"ℹ️ Week {year}-{month:02d}-W{week_num}: Exited as_completed loop. Accounted for {tasks_fully_accounted_for_this_week}/{total_tasks_for_week} tasks.")

                # This message prints just before the 'with' block for ProcessPoolExecutor ends
                tqdm.write(f"ℹ️ Week {year}-{month:02d}-W{week_num}: Attempting to exit ProcessPoolExecutor context (shutdown workers)...")
                # The 'with' block for ProcessPoolExecutor ends here. If workers are stuck, it might hang here.

                # This message prints AFTER the 'with' block has successfully exited
                tqdm.write(f"✅ Week {year}-{month:02d}-W{week_num}: Successfully exited ProcessPoolExecutor context.")
                # After the loop (either completed all tasks or failsafe triggered)
                if apply_failsafe_check and tasks_fully_accounted_for_this_week < total_tasks_for_week and tasks_fully_accounted_for_this_week < failsafe_tasks_processed_target:
                    # This means the loop exited, but not because of the failsafe, and not all tasks were processed.
                    # This could happen if `as_completed` stops yielding for some reason (e.g., BrokenProcessPool).
                    tqdm.write(f"⚠️ Note: Loop for week {year}-{month:02d}-W{week_num} ended. "
                               f"{tasks_fully_accounted_for_this_week}/{total_tasks_for_week} tasks accounted for. "
                               "This might indicate an issue if not all tasks were processed and failsafe didn't trigger as expected.")


                total_papers_overall += papers_this_week_processed_successfully
                # ... (rest of your weekly summary and CSV saving logic remains the same) ...
                if papers_this_week_processed_successfully > 0:
                    avg_task_time_for_week = task_time_week_successful / papers_this_week_processed_successfully
                    tqdm.write(f"📊 Weekly average processing time for {papers_this_week_processed_successfully} successful papers: {avg_task_time_for_week:.2f}s")
                if papers_this_week_timed_out > 0:
                    tqdm.write(f"⏰ {papers_this_week_timed_out} papers officially timed out this week.")
                if papers_this_week_other_exceptions > 0:
                    tqdm.write(f"🚫 {papers_this_week_other_exceptions} papers had other exceptions this week.")
                     
                 # --- Corrected CSV Saving Logic ---
                if master_primers_week: # Only proceed if there's data in master_primers_week
                    print(f"\n--- Consolidating Results for {year}-{month:02d}-Week{week_num} ---")
                    df_week = pd.DataFrame(master_primers_week) # Define df_week HERE

                    if not df_week.empty: # Further check if the created DataFrame is not empty
                        
                        # Define the full set of desired columns for consistency
                        cols_to_ensure = ['PMCID','Gene','Sequence','Orientation','Source File','Page','Original Cell Text','Citation Count']
                        
                        # Ensure all required columns exist, adding them with None if they are missing
                        for col in cols_to_ensure:
                            if col not in df_week.columns:
                                df_week[col] = None # Use None, or pd.NA for pandas specific NA
                        
                        # Select/reorder columns to the desired final set
                        df_week = df_week[cols_to_ensure]
                        
                        # Define columns for identifying duplicates
                        # Make sure these columns actually exist in df_week to avoid KeyErrors during subset
                        subset_dups_candidates = ['PMCID','Gene','Sequence','Orientation','Source File','Page']
                        subset_dups = [col for col in subset_dups_candidates if col in df_week.columns]
                        
                        if subset_dups: # Only drop duplicates if there are valid columns to use for subset
                            df_week.drop_duplicates(subset=subset_dups, inplace=True, keep='first')
                        else:
                            # This case might occur if essential key columns for deduplication are missing from master_primers_week items
                            print(f"⚠️ Could not perform deduplication for week {year}-{month:02d}-W{week_num} as key columns for subset were missing from DataFrame.")

                        num_unique_week = len(df_week)

                        if num_unique_week > 0: # Check again if DataFrame is non-empty after deduplication
                            total_primers_overall += num_unique_week
                            print(f"Unique primers for {year}-{month:02d}-Week{week_num}: {num_unique_week}")
                            try:
                                # Ensure OUT_DIR is defined (e.g., OUT_DIR = Path("./final_csvs_failsafe"))
                                csv_path = OUT_DIR / f"primers_{year}_{month:02d}_week_{week_num}.csv"
                                df_week.to_csv(csv_path, index=False, encoding='utf-8')
                                print(f"  ✅ Saved {num_unique_week} primers to {csv_path}")
                                shutil.rmtree("data/")
                                if os.path.exists("script_temp_files/"):
                                    for item in os.listdir("script_temp_files/"):
                                        item_path = os.path.join("script_temp_files/", item)
                                        if os.path.isfile(item_path):
                                            os.remove(item_path)
                                        elif os.path.isdir(item_path):
                                            shutil.rmtree(item_path)
                                print(f"  ✅ Saved {num_unique_week} primers to {csv_path}")
                            except Exception as e:
                                print(f"  ❌ Error saving CSV for week {year}-{month:02d}-W{week_num}: {e}")
                        else:
                             print(f"No unique primer data to save for week {year}-{month:02d}-W{week_num} (DataFrame became empty after deduplication).")
                    else:
                        # This case is if master_primers_week was not empty, but pd.DataFrame(master_primers_week) resulted in an empty df_week
                        print(f"No primer data to save for week {year}-{month:02d}-W{week_num} (DataFrame was empty after creation from master_primers_week).")
                else:
                    # This path is taken if master_primers_week list was empty from the start.
                    # df_week is NOT DEFINED in this path, and it's NOT USED.
                    print(f"No primers extracted for {year}-{month:02d}-Week{week_num} (master_primers_week list was empty).")

                batch_duration = time.time() - batch_start_time
                print(f"🏁 FINISHED {year}-{month:02d}-Week{week_num} in {batch_duration:.2f}s ({batch_duration/60:.2f}m) 🏁")


    total_run_time = time.time() - overall_start_time
    print(f"\n\n{'='*70}\n🎉 Batch processing COMPLETED. 🎉")
    print(f"Total papers processed successfully: {total_papers_overall}")
    print(f"Total papers officially timed out: {total_papers_timed_out_overall}")
    print(f"Total papers with other exceptions: {total_papers_other_exceptions_overall}")
    print(f"Total unique primers extracted: {total_primers_overall}")
    print(f"Total execution time: {total_run_time:.2f}s ({total_run_time/3600:.2f}hrs).\n{'='*70}")

    print(f"🧹 Final cleanup of script temporary directory: {TMP_DIR_FOR_SCRIPT}")
    remove_directory_recursively(TMP_DIR_FOR_SCRIPT)
    print("ℹ️ Script finished.")
